from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

def set_run_font(run, font_name="Times New Roman", font_size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def clear_cell(cell):
    cell.text = ""
    for p in cell.paragraphs:
        p.clear()

def set_cell_text(cell, text, align="left", valign="top", font_size=12, bold=False):
    clear_cell(cell)
    if valign == "center":
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    elif valign == "bottom":
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    else:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    lines = str(text or "").splitlines()
    if not lines:
        return

    for idx, line in enumerate(lines):
        if idx > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, font_size=font_size, bold=bold)

def fill_hang_hoa_table(docx_path, danh_sach_hang_hoa, action="thu_moi"):
    doc = Document(docx_path)
    target_table = None

    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if "STT" in table_text and "TÊN HÀNG HÓA" in table_text:
            target_table = table
            break

    if target_table is None:
        raise Exception("Không tìm thấy bảng hàng hóa trong file Word")

    sample_row_index = 1
    if len(target_table.rows) <= sample_row_index:
        raise Exception("Bảng hàng hóa chưa có dòng mẫu để copy style")

    sample_row = target_table.rows[sample_row_index]
    sample_cell_tcprs = [deepcopy(cell._tc.get_or_add_tcPr()) for cell in sample_row.cells]
    sample_defaults = [cell.text.strip() for cell in sample_row.cells]
    
    target_table._tbl.remove(sample_row._tr)

    if not danh_sach_hang_hoa:
        danh_sach_hang_hoa = []

    for item in danh_sach_hang_hoa:
        print("DEBUG DATA HÀNG HÓA NHẬN ĐƯỢC:", item, flush=True)
        row = target_table.add_row()
        for i, cell in enumerate(row.cells):
            if i < len(sample_cell_tcprs):
                old_tcPr = cell._tc.get_or_add_tcPr()
                cell._tc.remove(old_tcPr)
                cell._tc.insert(0, deepcopy(sample_cell_tcprs[i]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            
        set_cell_text(row.cells[0], str(item.get("STT", "")), align="center", valign="center")
        set_cell_text(row.cells[2], str(item.get("DVT", "")), align="center", valign="center")
        set_cell_text(row.cells[3], str(item.get("SoLuong", "")), align="center", valign="center")

        cell_ten = row.cells[1]
        clear_cell(cell_ten)
        p_ten = cell_ten.paragraphs[0]
        p_ten.paragraph_format.space_before = Pt(0)
        p_ten.paragraph_format.space_after = Pt(0)
        p_ten.paragraph_format.line_spacing = 1.0

        run_ten = p_ten.add_run(str(item.get("TenHang", "")))
        set_run_font(run_ten, bold=True)

        chi_tiet_render = str(item.get("ChiTietRender", ""))
        if chi_tiet_render:
            lines = chi_tiet_render.splitlines()
            for line in lines:
                p_ten.add_run().add_break()
                run_ct = p_ten.add_run(line)
                set_run_font(run_ct, bold=False)

        if action == "ban_giao":
            val_quy_cach = str(item.get("QuyCach", "")).strip() or sample_defaults[4]
            val_tinh_trang = str(item.get("TinhTrang", "")).strip() or sample_defaults[5]
            set_cell_text(row.cells[4], val_quy_cach, align="center", valign="center")
            set_cell_text(row.cells[5], val_tinh_trang, align="left", valign="center")
        elif action == "thanh_toan":
            set_cell_text(row.cells[4], str(item.get("DonGia", "")), align="right", valign="center")
            set_cell_text(row.cells[5], str(item.get("ThanhTien", "")), align="right", valign="center")
            set_cell_text(row.cells[6], str(item.get("chung_chi", "") or item.get("ChungChi", "")), align="left", valign="center")
        else: 
            set_cell_text(row.cells[4], str(item.get("chung_chi", "") or item.get("ChungChi", "")), align="left", valign="center")

    doc.save(docx_path)